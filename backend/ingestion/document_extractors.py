from __future__ import annotations

import os
import zipfile
from io import BytesIO
from typing import Any

from docx import Document
from lxml import etree

from backend.ingestion.models import ExtractedDocument, ExtractedPage, ImageAsset
from backend.ingestion.ocr_assets import OcrAssetProcessor
from backend.ingestion.pdf import PdfDocumentExtractor


class DocumentExtractor:
    def __init__(
        self,
        *,
        config: Any,
        trace_logger,
        run_ocr,
        ocr_worker_count,
        current_document_workers,
        detected_cpu_count,
        reserved_core_count,
        pdf_ext: str,
    ):
        self.config = config
        self.trace_logger = trace_logger
        self.run_ocr = run_ocr
        self.ocr_worker_count = ocr_worker_count
        self.current_document_workers = current_document_workers
        self.detected_cpu_count = detected_cpu_count
        self.reserved_core_count = reserved_core_count
        self.pdf_ext = pdf_ext

    def extract_by_type(self, fpath: str, chunker, progress_callback=None) -> ExtractedDocument | None:
        if os.path.isdir(fpath):
            self.trace_logger.warning(f"Skipping directory: {fpath}")
            return None

        ext = os.path.splitext(fpath)[1].lower()
        if ext == self.pdf_ext:
            return self.extract_pdf(fpath, chunker, progress_callback)
        if ext == ".docx":
            return self.extract_docx(fpath, chunker)
        if ext in {".md", ".txt"}:
            return self.extract_text(fpath)
        if ext in {".png", ".jpg", ".jpeg"}:
            return self.extract_image(fpath, chunker)

        self.trace_logger.warning(f"Unsupported file type: {ext} for file: {fpath}")
        return None

    def extract_pdf(self, fpath: str, chunker, progress_callback=None) -> ExtractedDocument:
        return PdfDocumentExtractor(
            config=self.config,
            ocr_assets=self._ocr_assets(chunker),
            trace_logger=self.trace_logger,
        ).extract(fpath, progress_callback=progress_callback)

    def extract_docx(self, fpath: str, chunker) -> ExtractedDocument:
        doc = Document(fpath)
        text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        document = ExtractedDocument(source_path=fpath, pages=[ExtractedPage(page_number=1, text=text)])
        document.assets.extend(self.extract_docx_textbox_images(fpath, chunker))
        return document

    def extract_text(self, fpath: str) -> ExtractedDocument:
        with open(fpath, "r", encoding="utf-8") as file:
            text = file.read()
        return ExtractedDocument(source_path=fpath, pages=[ExtractedPage(page_number=1, text=text)])

    def extract_image(self, fpath: str, chunker) -> ExtractedDocument:
        with open(fpath, "rb") as image_file:
            raw = image_file.read()
        result = self._ocr_assets(chunker).run_jobs([(0, 1, raw, os.path.basename(fpath), "image")])[0]
        asset = PdfDocumentExtractor._asset_from_ocr_result(result, os.path.basename(fpath), {})
        return ExtractedDocument(
            source_path=fpath,
            pages=[ExtractedPage(page_number=1, text=asset.ocr_text)],
            assets=[asset],
        )

    def extract_docx_textbox_images(self, fpath: str, chunker) -> list[ImageAsset]:
        base_doc = os.path.basename(fpath)
        assets: list[ImageAsset] = []
        try:
            image_jobs = self._docx_textbox_image_jobs(fpath, base_doc)
            for result in self._ocr_assets(chunker).run_jobs(image_jobs):
                ocr_result = result["ocr_result"]
                if not ocr_result["accepted"]:
                    continue
                assets.append(
                    ImageAsset(
                        image_key=result["image_key"],
                        page_number=1,
                        caption=f"Textbox image in {base_doc}",
                        image_bytes=result["image_bytes"],
                        searchable_text=ocr_result["text"],
                        ocr_text=ocr_result["text"],
                        ocr_score=float(ocr_result.get("score") or 0.0),
                        ocr_confidence=ocr_result.get("confidence"),
                        source_kind="textbox",
                    )
                )
        except Exception as exc:
            self.trace_logger.warning(f"Failed to extract DOCX textbox images from {fpath}: {exc}")
        return assets

    def _ocr_assets(self, chunker) -> OcrAssetProcessor:
        return OcrAssetProcessor(
            config=self.config,
            chunker=chunker,
            run_ocr=self.run_ocr,
            trace_logger=self.trace_logger,
            ocr_worker_count=self.ocr_worker_count,
            current_document_workers=self.current_document_workers,
            detected_cpu_count=self.detected_cpu_count,
            reserved_core_count=self.reserved_core_count,
        )

    def _docx_textbox_image_jobs(self, fpath: str, base_doc: str) -> list[tuple]:
        with open(fpath, "rb") as file:
            docx_content = file.read()
        with zipfile.ZipFile(BytesIO(docx_content)) as docx_zip:
            rels = self._docx_relationships(docx_zip)
            if "word/document.xml" not in docx_zip.namelist():
                return []
            tree = etree.fromstring(docx_zip.read("word/document.xml"))
            namespaces = {
                "v": "urn:schemas-microsoft-com:vml",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            }
            image_jobs = []
            for idx, v_img in enumerate(tree.xpath(".//v:imagedata", namespaces=namespaces)):
                r_embed = v_img.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                if not r_embed or r_embed not in rels:
                    continue
                img_path = "word/" + rels[r_embed].replace("\\", "/")
                if img_path not in docx_zip.namelist():
                    continue
                image_jobs.append((len(image_jobs), 1, docx_zip.read(img_path), f"{base_doc}_textbox_img{idx+1}", "textbox"))
            return image_jobs

    @staticmethod
    def _docx_relationships(docx_zip: zipfile.ZipFile) -> dict[str, str]:
        rels = {}
        for name in docx_zip.namelist():
            if not name.startswith("word/_rels/") or not name.endswith(".xml.rels"):
                continue
            rel_tree = etree.fromstring(docx_zip.read(name))
            for rel in rel_tree.xpath(
                "//rel:Relationship",
                namespaces={"rel": "http://schemas.openxmlformats.org/package/2006/relationships"},
            ):
                rid = rel.attrib["Id"]
                target = rel.attrib["Target"]
                rels[rid] = os.path.normpath(os.path.join(os.path.dirname(name), target))
        return rels
