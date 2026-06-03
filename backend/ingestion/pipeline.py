from __future__ import annotations

import os
import re
import zipfile
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from typing import Any

from docx import Document
from lxml import etree

from backend.ingestion.document_classifier import classify_document
from backend.ingestion.models import ExtractedDocument, ExtractedPage, ImageAsset
from backend.ingestion.ocr_assets import OcrAssetProcessor
from backend.ingestion.pdf_extractor import PdfExtractor
from pdf_visuals import link_chunks_to_rendered_pages


class IngestionPipeline:
    def __init__(
        self,
        *,
        config: Any,
        trace_logger,
        run_ocr,
        detected_cpu_count,
        reserved_core_count,
        usable_core_count,
        document_worker_count,
        ocr_worker_count,
        current_document_workers,
        begin_document_worker,
        finish_document_worker,
        pdf_ext: str,
    ):
        self.config = config
        self.trace_logger = trace_logger
        self.run_ocr = run_ocr
        self.detected_cpu_count = detected_cpu_count
        self.reserved_core_count = reserved_core_count
        self.usable_core_count = usable_core_count
        self.document_worker_count = document_worker_count
        self.ocr_worker_count = ocr_worker_count
        self.current_document_workers = current_document_workers
        self.begin_document_worker = begin_document_worker
        self.finish_document_worker = finish_document_worker
        self.pdf_ext = pdf_ext

    def process_file_by_type(self, fpath, target_state, trace_logger, chunker, token_utils, progress_callback=None):
        if os.path.isdir(fpath):
            trace_logger.warning(f"Skipping directory: {fpath}")
            return

        ext = os.path.splitext(fpath)[1].lower()
        if ext == self.pdf_ext:
            document = self._extract_pdf(fpath, chunker, progress_callback)
        elif ext == ".docx":
            document = self._extract_docx(fpath, chunker)
        elif ext in {".md", ".txt"}:
            document = self._extract_text(fpath)
        elif ext in {".png", ".jpg", ".jpeg"}:
            document = self._extract_image(fpath, chunker)
        else:
            trace_logger.warning(f"Unsupported file type: {ext} for file: {fpath}")
            return

        self._apply_profile(document)
        if progress_callback:
            progress_callback(
                currentDocument=os.path.basename(fpath),
                documentPhase="Chunking extracted text",
                documentType=document.profile.document_type if document.profile else "unknown",
            )
        self._store_extracted_document(document, target_state, chunker, token_utils)

    def load_documents_parallel(
        self,
        folder,
        files_selected,
        clear_existing=True,
        target_state=None,
        target_chunker=None,
        target_token_utils=None,
        progress_callback=None,
    ):
        if target_state is None or target_chunker is None or target_token_utils is None:
            raise ValueError("target_state, target_chunker, and target_token_utils are required.")
        if clear_existing:
            target_state.clear_all()

        if isinstance(files_selected, list):
            file_list = list(files_selected)
        else:
            file_list = self._scan_folder(folder)
        file_list.sort(key=_extract_first_number)

        if progress_callback:
            progress_callback(stage="processing_documents", total_files=len(file_list))
        if not file_list:
            self.trace_logger.warning(f"No supported documents found in {folder}.")
            return

        def process_file(filename):
            fpath = filename if os.path.isabs(filename) else os.path.join(folder, filename)
            active_count = self.begin_document_worker()
            try:
                if progress_callback:
                    progress_callback(stage="processing_documents", current_file=filename)
                thread_id = threading.get_ident()
                self.trace_logger.info(f"Thread-{thread_id} started for {filename} ({active_count} active document workers)")
                started = time.time()

                def detail_progress(**details):
                    if progress_callback:
                        progress_callback(stage="processing_documents", current_file=filename, file_details=details)

                self.process_file_by_type(
                    fpath,
                    target_state,
                    self.trace_logger,
                    target_chunker,
                    target_token_utils,
                    progress_callback=detail_progress,
                )
                self.trace_logger.info(f"Thread-{thread_id} finished {filename} in {time.time() - started:.2f}s")
            finally:
                if progress_callback:
                    progress_callback(stage="processing_documents", finished_file=filename)
                self.finish_document_worker()

        cpu_count = self.detected_cpu_count()
        max_workers = self.document_worker_count(len(file_list), cpu_count=cpu_count)
        self.trace_logger.info(
            f"Ingest worker budget: {cpu_count} cores detected, reserving {self.reserved_core_count(cpu_count)}, "
            f"{self.usable_core_count(cpu_count)} usable, {max_workers} document workers for {len(file_list)} files."
        )
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(process_file, filename): filename for filename in file_list}
            for future in as_completed(futures):
                future.result()

    def _scan_folder(self, folder: str) -> list[str]:
        if not os.path.exists(folder):
            self.trace_logger.error("Training folder not found.")
            return []
        recursive = self.config.get("TRAINING_RECURSIVE", True)
        excluded_dirs = set(self.config.get("TRAINING_EXCLUDE_DIRS", []))
        supported = {self.pdf_ext, ".docx", ".md", ".txt", ".png", ".jpg", ".jpeg"}
        file_list = []
        if recursive:
            for root, dirnames, filenames in os.walk(folder):
                dirnames[:] = [
                    dirname for dirname in dirnames
                    if dirname not in excluded_dirs
                    and os.path.relpath(os.path.join(root, dirname), folder) not in excluded_dirs
                ]
                for filename in filenames:
                    if filename.startswith("~$"):
                        continue
                    fpath = os.path.join(root, filename)
                    if os.path.isfile(fpath) and os.path.getsize(fpath) > 0 and os.path.splitext(filename)[1].lower() in supported:
                        file_list.append(os.path.relpath(fpath, folder))
            return file_list
        for filename in os.listdir(folder):
            fpath = os.path.join(folder, filename)
            if (
                not filename.startswith("~$")
                and os.path.isfile(fpath)
                and os.path.getsize(fpath) > 0
                and os.path.splitext(filename)[1].lower() in supported
            ):
                file_list.append(filename)
        return file_list

    def _extract_pdf(self, fpath: str, chunker, progress_callback=None) -> ExtractedDocument:
        ocr_assets = OcrAssetProcessor(
            config=self.config,
            chunker=chunker,
            run_ocr=self.run_ocr,
            trace_logger=self.trace_logger,
            ocr_worker_count=self.ocr_worker_count,
            current_document_workers=self.current_document_workers,
            detected_cpu_count=self.detected_cpu_count,
            reserved_core_count=self.reserved_core_count,
        )
        return PdfExtractor(config=self.config, ocr_assets=ocr_assets, trace_logger=self.trace_logger).extract(
            fpath,
            progress_callback=progress_callback,
        )

    def _extract_docx(self, fpath: str, chunker) -> ExtractedDocument:
        doc = Document(fpath)
        text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        document = ExtractedDocument(source_path=fpath, pages=[ExtractedPage(page_number=1, text=text)])
        document.assets.extend(self._extract_docx_textbox_images(fpath, chunker))
        return document

    def _extract_text(self, fpath: str) -> ExtractedDocument:
        with open(fpath, "r", encoding="utf-8") as file:
            text = file.read()
        return ExtractedDocument(source_path=fpath, pages=[ExtractedPage(page_number=1, text=text)])

    def _extract_image(self, fpath: str, chunker) -> ExtractedDocument:
        with open(fpath, "rb") as image_file:
            raw = image_file.read()
        ocr_assets = OcrAssetProcessor(
            config=self.config,
            chunker=chunker,
            run_ocr=self.run_ocr,
            trace_logger=self.trace_logger,
            ocr_worker_count=self.ocr_worker_count,
            current_document_workers=self.current_document_workers,
            detected_cpu_count=self.detected_cpu_count,
            reserved_core_count=self.reserved_core_count,
        )
        result = ocr_assets.run_jobs([(0, 1, raw, os.path.basename(fpath), "image")])[0]
        asset = PdfExtractor._asset_from_ocr_result(result, os.path.basename(fpath), {})
        return ExtractedDocument(
            source_path=fpath,
            pages=[ExtractedPage(page_number=1, text=asset.ocr_text)],
            assets=[asset],
        )

    def _apply_profile(self, document: ExtractedDocument) -> None:
        document.profile = classify_document(document.source_path, document.pages)
        profile = document.profile
        self.trace_logger.info(
            f"Ingest profile for {os.path.basename(document.source_path)}: "
            f"{profile.document_type} ({profile.confidence:.2f})"
            + (f", component {profile.component_name}" if profile.component_name else "")
        )

    def _store_extracted_document(self, document: ExtractedDocument, target_state, chunker, token_utils) -> None:
        page_texts = [page.text for page in document.pages]
        density = token_utils.estimate_token_density("\n\n".join(page_texts))
        if density > 40:
            self.trace_logger.warning(f"High token density in {document.source_path}: {density:.2f} tokens/line")

        chunks, metadata = chunker.smart_chunk_pages(page_texts, document.source_path)
        profile_meta = document.profile.metadata() if document.profile else {}
        for meta in metadata:
            meta.update(profile_meta)
            meta["parent_source"] = document.source_path

        rendered_image_pages = {
            asset.page_number: asset.image_key
            for asset in document.assets
            if asset.source_kind == "rendered" and asset.page_number
        }
        linked_visuals = link_chunks_to_rendered_pages(
            chunks,
            metadata,
            document.source_path,
            {asset.image_key for asset in document.assets},
        )
        for meta in metadata:
            page = _optional_int(meta.get("page"))
            if page in rendered_image_pages and not meta.get("source_image_id") and _chunk_mentions_visual(meta):
                meta["source_image_id"] = rendered_image_pages[page]
        if linked_visuals:
            self.trace_logger.info(f"Linked {linked_visuals} text chunks to rendered page images for {os.path.basename(document.source_path)}")

        image_chunks, image_sources, image_meta = [], [], []
        min_chars = int(self.config.get("OCR_INDEX_TEXT_MIN_CHARS", 80) or 80)
        for asset in document.assets:
            target_state.add_image_store(asset.image_key, OcrAssetProcessor.base64_image(asset.image_bytes))
            target_state.add_image_caption(asset.image_key, asset.caption)
            target_state.add_image_mime_type(asset.image_key, asset.mime_type)
            if asset.searchable_text:
                target_state.add_image_page_text(asset.image_key, asset.searchable_text)
            if self.config.get("INDEX_IMAGE_OCR_AS_TEXT", False) and asset.ocr_text and len(asset.ocr_text) >= min_chars:
                image_chunks.append(asset.ocr_text)
                image_sources.append(document.source_path)
                ocr_meta = chunker.make_chunk_meta(asset.ocr_text, document.source_path, self._ocr_section(asset), "ocr")
                ocr_meta.update(profile_meta)
                ocr_meta.update({
                    "page": asset.page_number,
                    "parent_source": document.source_path,
                    "source_image_id": asset.image_key,
                    "ocr_score": asset.ocr_score,
                    "ocr_confidence": asset.ocr_confidence,
                    "chunk_type": "ocr",
                })
                image_meta.append(ocr_meta)

        target_state.extend_chunks(
            chunks + image_chunks,
            [document.source_path] * len(chunks) + image_sources,
            metadata + image_meta,
        )

    @staticmethod
    def _ocr_section(asset: ImageAsset) -> str:
        return "Rendered Page OCR" if asset.source_kind == "rendered" else "Image OCR"

    def _extract_docx_textbox_images(self, fpath: str, chunker) -> list[ImageAsset]:
        base_doc = os.path.basename(fpath)
        assets: list[ImageAsset] = []
        try:
            with open(fpath, "rb") as file:
                docx_content = file.read()
            with zipfile.ZipFile(BytesIO(docx_content)) as docx_zip:
                rels = {}
                for name in docx_zip.namelist():
                    if name.startswith("word/_rels/") and name.endswith(".xml.rels"):
                        rel_tree = etree.fromstring(docx_zip.read(name))
                        for rel in rel_tree.xpath(
                            "//rel:Relationship",
                            namespaces={"rel": "http://schemas.openxmlformats.org/package/2006/relationships"},
                        ):
                            rid = rel.attrib["Id"]
                            target = rel.attrib["Target"]
                            rels[rid] = os.path.normpath(os.path.join(os.path.dirname(name), target))

                if "word/document.xml" not in docx_zip.namelist():
                    return assets
                tree = etree.fromstring(docx_zip.read("word/document.xml"))
                namespaces = {
                    "v": "urn:schemas-microsoft-com:vml",
                    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                }
                image_jobs = []
                for idx, v_img in enumerate(tree.xpath(".//v:imagedata", namespaces=namespaces)):
                    r_embed = v_img.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                    if not r_embed or r_embed not in rels:
                        continue
                    img_path = "word/" + rels[r_embed].replace("\\", "/")
                    if img_path not in docx_zip.namelist():
                        continue
                    image_jobs.append((len(image_jobs), 1, docx_zip.read(img_path), f"{base_doc}_textbox_img{idx+1}", "textbox"))

            ocr_assets = OcrAssetProcessor(
                config=self.config,
                chunker=chunker,
                run_ocr=self.run_ocr,
                trace_logger=self.trace_logger,
                ocr_worker_count=self.ocr_worker_count,
                current_document_workers=self.current_document_workers,
                detected_cpu_count=self.detected_cpu_count,
                reserved_core_count=self.reserved_core_count,
            )
            for result in ocr_assets.run_jobs(image_jobs):
                ocr_result = result["ocr_result"]
                if not ocr_result["accepted"]:
                    continue
                assets.append(
                    ImageAsset(
                        image_key=result["image_key"],
                        page_number=1,
                        caption=f"Textbox image in {base_doc}",
                        image_bytes=result["image_bytes"],
                        searchable_text=ocr_result["text"],
                        ocr_text=ocr_result["text"],
                        ocr_score=float(ocr_result.get("score") or 0.0),
                        ocr_confidence=ocr_result.get("confidence"),
                        source_kind="textbox",
                    )
                )
        except Exception as exc:
            self.trace_logger.warning(f"Failed to extract DOCX textbox images from {fpath}: {exc}")
        return assets

    @staticmethod
    def extract_page_number(value) -> int | None:
        match = re.search(r"_page(\d+)", str(value or ""))
        return int(match.group(1)) if match else None


def _optional_int(value) -> int | None:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return None
    return number if number > 0 else None


def _extract_first_number(value: str) -> int:
    match = re.search(r"(\d+)", value)
    return int(match.group(1)) if match else 0


def _chunk_mentions_visual(meta: dict) -> bool:
    text = " ".join(str(meta.get(key) or "") for key in ("section", "category", "chunk_type", "visual_references"))
    return bool(re.search(r"\b(fig(?:ure)?|diagram|schematic|pinout|layout|image|rendered|ocr)\b", text, re.IGNORECASE))
